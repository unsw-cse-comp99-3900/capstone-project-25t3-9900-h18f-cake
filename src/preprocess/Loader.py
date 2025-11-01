'''data loader+ metadata extractor
1. load file
2. transfer to UTF-8(PDF/OCR/WORD)
2. extract metadata
'''
import os
import zipfile
from typing import Optional
import fitz
from paddleocr import PaddleOCR
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from tqdm import tqdm
import paddle
import numpy as np
import cv2
paddle.set_device("cpu")
# OCR 依赖（可选）
try:
    import fitz  # PyMuPDF
    from PIL import Image
    import pytesseract
    _OCR_AVAILABLE = True
except Exception:
    _OCR_AVAILABLE = False


class DataLoader:
    def __init__(
        self,
        ocr_language: str = "en",
        min_chars: int = 20,
        alpha_ratio: float = 0.05,
        dpi: int = 72,#OCR 渲染清晰度【尝试过了 足够了 再大提取不出来了】
        min_conf = 0.7,
        max_pages: Optional[int] = None,
    ):

        self.ocr_lang = ocr_language
        self.min_chars = min_chars
        self.alpha_ratio = alpha_ratio
        self.dpi = dpi
        self.max_pages = max_pages
        self.min_conf = min_conf

    def utf8_normalize(self, text):
        if isinstance(text, str):
            data = text.strip()
            return data.encode("utf-8", errors="ignore").decode("utf-8")
        elif isinstance(text, list):
            return [self.utf8_normalize(item) for item in text]
        elif isinstance(text, dict):
            return {k: self.utf8_normalize(v) for k, v in text.items()}
        else:
            return text
    def table2text(self, table):
        rows = []
        for r in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            rows.append(cells)
        if not rows:
            return ""
    
        header = "| " + " | ".join(rows[0]) + " |"
        sep    = "| " + " | ".join(["---"] * len(rows[0])) + " |"
        body   = "\n".join("| " + " | ".join(row) + " |" for row in rows[1:])
        return "\n".join([header, sep, body])
    def image2bytes(self, b, out_dir, stem, idx, ext) :
        os.makedirs(out_dir, exist_ok=True)
        fname = f"{stem}_{idx}{ext}"
        fpath = os.path.join(out_dir, fname)
        with open(fpath, "wb") as f:
            f.write(b)
        return fpath
    def load_docx(self, path,image_dir ):
        try:
            file = Document(path)
            # import pdb
            # pdb.set_trace()   
        except Exception:
            return {"paragraphs": [], "tables": [], "images": [], "meta": {}}
        paragraphs, tables, images = [], [], []
        for idx, para in enumerate(file.paragraphs):
            txt = (para.text or "").strip()
            if txt:
                paragraphs.append({"para_id": idx, "text": txt})
        for idx, tab in enumerate(file.tables, start=1):
            table_txt = self.table2text(tab)
            if table_txt.strip():
                tables.append({"table_id": idx, "markdown": table_txt})
        img_idx = 1
        stem = os.path.splitext(os.path.basename(path))[0]
        if image_dir:
            out_dir = os.path.join(image_dir, stem)
            os.makedirs(out_dir, exist_ok=True)

            with zipfile.ZipFile(path, "r") as docx_zip:
                media_files = [f for f in docx_zip.namelist() if f.startswith("word/media/")]
                for i, filename in enumerate(media_files, start=1):
                    # print(filename)
                    ext = os.path.splitext(filename)[1]
                    img_data = docx_zip.read(filename)
                    img_path = os.path.join(out_dir, f"img_{i}{ext}")
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    images.append({
                        "image_id": i,
                        "path": img_path,
                        "caption": "",
                        "source": "archive_extract"
                    })

            # 匹配 figure caption（可选）
            para_texts = [p.text.strip() for p in file.paragraphs]
            for img in images:
                for txt in para_texts:
                    if txt.lower().startswith(("figure", "fig.")) and str(img["image_id"]) in txt:
                        img["caption"] = txt
                        break

        print(f"[INFO] Extracted {len(images)} images from {path}")

        meta = {"zid":os.path.basename(path).split('.')[0],"ext": os.path.basename(path).split('.')[1], "filename": os.path.basename(path)}
        return {"paragraphs": paragraphs, "tables": tables, "images": images, "meta": meta}

        
    def load_pdf(self, path): 
        try:
            doc = fitz.open(path)
            parts = [page.get_text("text") for page in doc]
            text = "\n".join(parts).strip()
            return text
        except Exception:
            return ""
        
    def file_scanned(self, text): #whether scanned
        if not text or len(text) < self.min_chars:
            return True
        readable = sum(ch.isalnum() or ("\u4e00" <= ch <= "\u9fff") for ch in text)
        return (readable / max(len(text), 1)) < self.alpha_ratio

    def ocr_pdf(self, path):
        ocr = getattr(self, "ocr", None)
        if ocr is None:
            ocr = PaddleOCR(use_textline_orientation=True, lang="en")
            try:
                paddle.set_device("cpu")
            except Exception:
                pass
            self.ocr = ocr
        doc = fitz.open(path)
        texts = []
        n = len(doc)
        mat = fitz.Matrix(self.dpi / 72, self.dpi / 72)

        total_pages = n if getattr(self, "max_pages", None) is None else min(n, self.max_pages)

        for page_idx in tqdm(range(total_pages), desc="OCR Pages", unit="page"):
            page = doc[page_idx]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            if pix.n == 4:
                img = cv2.cvtColor(img, cv2.COLOR_RGBA2BGR)
            elif pix.n == 3:
                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
            elif pix.n == 1:
                img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
            # cv2.imwrite(f"debug_page_{page_idx+1}.png", img)
            # print('1',img.shape, img.dtype)
            
            result = ocr.predict(img)
            page_lines = []

            if isinstance(result, dict):
                rec_texts  = result.get("rec_texts", []) or []
                rec_scores = result.get("rec_scores", []) or []
                if self.min_conf > 0:
                    for t, sc in zip(rec_texts, rec_scores or [1.0] * len(rec_texts)):
                        if sc is None or sc >= self.min_conf:
                            page_lines.append(t)
                else:
                    page_lines = rec_texts

            elif isinstance(result, list) and result:
                if isinstance(result[0], dict):
                    payload = result[0]
                    rec_texts  = payload.get("rec_texts", []) or []
                    rec_scores = payload.get("rec_scores", []) or []
                    # print(rec_scores)
                    if self.min_conf > 0:
                        for t, sc in zip(rec_texts, rec_scores or [1.0] * len(rec_texts)):
                            if sc is None or sc >= self.min_conf:
                                page_lines.append(t)
                    else:
                        page_lines = rec_texts
                else:
                    lines = result[0] if isinstance(result[0], list) else result
                    for seg in lines:
                        try:
                            _, (txt, conf) = seg
                        except Exception:
                            continue
                        if not txt:
                            continue
                        if self.min_conf and conf is not None and conf < self.min_conf:
                            continue
                        page_lines.append(txt)
            texts.append("\n".join(page_lines).strip())

        doc.close()
        return "\n".join(texts).strip()

    def load_file(self, path, image_dir):
        file_form = os.path.splitext(path)[1].lower()
        if file_form == ".docx" or file_form == ".doc":
            file_txt = self.load_docx(path,image_dir)
            return self.utf8_normalize(file_txt)
        elif file_form == ".pdf":
            file_txt = self.load_pdf(path)
            use_ocr = self.file_scanned(file_txt)
            # print('\t[USE_OCR]',use_ocr)
            if use_ocr:
                file_txt = self.ocr_pdf(path)
            # print(file_txt)
            return self.utf8_normalize(file_txt)
        else:
            print(f"[WARN] WRONG FILE FORM:{path})")
         
    def metadata_extraction(self, filepath: str):

        stem, _ = os.path.splitext(os.path.basename(filepath))
        parts = stem.split("_")
        
        if len(parts) < 3:
            return None, None, None  

        if not (parts[1] and parts[1].lower().startswith("z") and parts[1][1:].isdigit()):
            parts[1] = None
        #assignment_id, student_id, grader
        return parts[0], parts[1], parts[2]







# if __name__ == "__main__":
#     from pprint import pprint

#     loader = DataLoader()

#     sample_path = "/Users/chenjo/Desktop/UNSW/2025/9900/AI_Moule/data/raw/z111.docx"
#     image_dir = "/Users/chenjo/Desktop/UNSW/2025/9900/AI_Moule/data/raw"
#     if os.path.exists(sample_path):
#         text_recd = loader.load_file(sample_path,image_dir)
#         # for i in text_recd['paragraphs']:
#         #     print(i['para_id'],'\n',i['text'])
#         # for i in text_recd['tables']:
#         #     print(i['markdown'])
#         print(text_recd)
 
#     else:
#         print(f"[WARNing]")